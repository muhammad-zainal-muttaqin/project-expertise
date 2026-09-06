"""Render laporan review ke DOCX dan periksa struktur/rujukan lokal."""
from __future__ import annotations

import json
import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.opc.constants import RELATIONSHIP_TYPE as RT

HERE = Path(__file__).resolve().parent
SOURCE = HERE / 'report-source.md'
OUTPUT = HERE / 'review-project-expertise-2026-09-06.docx'
LINK = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')


def inline(paragraph, text, links):
    offset = 0
    for match in LINK.finditer(text):
        paragraph.add_run(text[offset:match.start()].replace('`', '').replace('**', ''))
        label, target = match.groups()
        if not target.startswith('https://'):
            resolved = (HERE / target).resolve()
            if not resolved.is_file():
                raise FileNotFoundError(resolved)
            target = str(resolved)
        relation = paragraph.part.relate_to(target, RT.HYPERLINK, is_external=True)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), relation)
        run = OxmlElement('w:r')
        prop = OxmlElement('w:rPr')
        color = OxmlElement('w:color'); color.set(qn('w:val'), '17665D')
        underline = OxmlElement('w:u'); underline.set(qn('w:val'), 'single')
        prop.extend([color, underline]); run.append(prop)
        content = OxmlElement('w:t'); content.text = label
        run.append(content); hyperlink.append(run); paragraph._p.append(hyperlink)
        links.append({'label': label, 'target': target})
        offset = match.end()
    paragraph.add_run(text[offset:].replace('`', '').replace('**', ''))


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    section.top_margin = section.bottom_margin = Cm(1.8)
    section.left_margin = section.right_margin = Cm(1.8)
    normal = doc.styles['Normal']
    normal.font.name, normal.font.size = 'Calibri', Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08
    for name, size in [('Title', 23), ('Heading 1', 16), ('Heading 2', 12)]:
        style = doc.styles[name]
        style.font.name, style.font.size = 'Calibri', Pt(size)
        style.font.color.rgb = RGBColor.from_string('174B43')
        style.paragraph_format.keep_with_next = True
    header = section.header.paragraphs[0]
    header.text = 'PROJECT-EXPERTISE  /  REVIEW IMPLEMENTASI'
    header.runs[0].font.size = Pt(8)
    footer = section.footer.paragraphs[0]
    footer.add_run('6 September 2026  ·  Diagnosis, tanpa pelatihan baru  |  ')
    field = OxmlElement('w:fldSimple'); field.set(qn('w:instr'), 'PAGE')
    footer._p.append(field)
    for run in footer.runs:
        run.font.size = Pt(8)
    lines = SOURCE.read_text().splitlines()
    links, i, tables = [], 0, 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1; continue
        if line.startswith('|'):
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                if not all(re.fullmatch(r'[:\- ]+', c) for c in cells):
                    rows.append(cells)
                i += 1
            table = doc.add_table(rows=0, cols=len(rows[0]))
            table.style = 'Light Shading Accent 1'
            table.autofit = False
            fractions = [.45, .275, .275] if rows[0][0] == 'Besaran' else [.20, .40, .40]
            for row_index, cells in enumerate(rows):
                row = table.add_row()
                properties = row._tr.get_or_add_trPr()
                properties.append(OxmlElement('w:cantSplit'))
                if row_index == 0:
                    properties.append(OxmlElement('w:tblHeader'))
                for column, value in enumerate(cells):
                    cell = row.cells[column]; cell.width = Cm(17.4 * fractions[column])
                    inline(cell.paragraphs[0], value, links)
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(9)
                        run.bold = row_index == 0
            doc.add_paragraph()
            tables += 1
            continue
        if line.startswith('# '):
            paragraph = doc.add_paragraph(style='Title'); line = line[2:]
        elif line.startswith('## '):
            paragraph = doc.add_paragraph(style='Heading 1'); line = line[3:]
        elif line.startswith('### '):
            paragraph = doc.add_paragraph(style='Heading 2'); line = line[4:]
        elif re.match(r'\d+\. ', line):
            paragraph = doc.add_paragraph(style='List Number')
            line = re.sub(r'^\d+\. ', '', line)
        else:
            paragraph = doc.add_paragraph()
        inline(paragraph, line, links)
        i += 1
    doc.core_properties.title = 'Review implementasi dan validitas eksperimen project-expertise'
    doc.core_properties.subject = 'Audit kode, metrik, dan generalisasi; AF-E-001 sampai AF-E-013'
    doc.core_properties.author = 'Codex — review untuk pengembang project-expertise'
    doc.save(OUTPUT)
    reopened = Document(OUTPUT)
    with ZipFile(OUTPUT) as archive:
        assert archive.testzip() is None
        xml = archive.read('word/document.xml').decode()
    assert len(reopened.tables) == tables == 2
    assert xml.count('<w:hyperlink ') == len(links)
    for expected in ['0,5201', '0,5198', '0,6569', '92,91%', '39 pohon']:
        assert expected in xml, expected
    qa = {'artifact': str(OUTPUT), 'bytes': OUTPUT.stat().st_size,
          'paragraphs': len(reopened.paragraphs), 'tables': tables,
          'hyperlinks': len(links), 'local_links_verified': sum(not x['target'].startswith('https://') for x in links),
          'visual_review': 'unavailable: no document renderer installed',
          'structural_review': 'passed'}
    (HERE / '../../results/audit_2026-09-06/report_qa.json').resolve().write_text(json.dumps(qa, indent=2) + '\n')
    print(json.dumps(qa, indent=2))


if __name__ == '__main__':
    main()
